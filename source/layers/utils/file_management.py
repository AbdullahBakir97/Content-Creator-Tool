import os
import shutil
import magic
from pathlib import Path
from apps.core.services import BaseService
from django.core.exceptions import ValidationError
from typing import Optional, List, Dict, Any, Union, BinaryIO, Tuple
import logging
from datetime import datetime, timedelta
import hashlib
from PIL import Image
import json
from django.conf import settings
import tempfile
import mimetypes
import zipfile
import PyPDF2
from docx import Document
from docx.shared import Inches
import fitz  # PyMuPDF

logger = logging.getLogger(__name__)

class FileManagementUtility(BaseService):
    """Utility for managing files with advanced features"""

    ALLOWED_IMAGE_TYPES = {'image/jpeg', 'image/png', 'image/gif', 'image/webp'}
    ALLOWED_VIDEO_TYPES = {'video/mp4', 'video/quicktime', 'video/x-msvideo'}
    ALLOWED_AUDIO_TYPES = {'audio/mpeg', 'audio/wav', 'audio/ogg'}
    ALLOWED_DOCUMENT_TYPES = {
        'application/pdf',
        'application/msword',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'text/plain',
        'application/rtf'
    }

    DOCUMENT_EXTENSIONS = {'.txt', '.pdf', '.doc', '.docx', '.rtf'}

    def __init__(self):
        super().__init__()
        self.temp_dir = os.path.join(settings.MEDIA_ROOT, 'temp')
        self.media_dir = settings.MEDIA_ROOT
        os.makedirs(self.temp_dir, exist_ok=True)

    def _validate(self, data: Dict[str, Any]) -> None:
        """Validate input data"""
        if 'file_path' not in data and 'file_content' not in data:
            raise ValidationError("Either file path or content is required")

    def get_mime_type(self, file_path: str) -> Optional[str]:
        """Get MIME type of file"""
        try:
            return magic.from_file(file_path, mime=True)
        except Exception as e:
            self.add_error(f"Failed to get MIME type: {str(e)}")
            return None

    def validate_file_type(self, file_path: str, allowed_types: set) -> bool:
        """Validate file type against allowed types"""
        mime_type = self.get_mime_type(file_path)
        return mime_type in allowed_types if mime_type else False

    def generate_unique_filename(self, original_name: str, directory: str) -> str:
        """Generate unique filename using hash"""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        name, ext = os.path.splitext(original_name)
        hash_input = f"{name}{timestamp}{os.urandom(8).hex()}"
        filename_hash = hashlib.md5(hash_input.encode()).hexdigest()[:8]
        return f"{name}_{timestamp}_{filename_hash}{ext}"

    def create_directory_structure(self, base_dir: str, structure: Dict) -> None:
        """Create nested directory structure"""
        try:
            for name, content in structure.items():
                path = os.path.join(base_dir, name)
                if isinstance(content, dict):
                    os.makedirs(path, exist_ok=True)
                    self.create_directory_structure(path, content)
                else:
                    os.makedirs(os.path.dirname(path), exist_ok=True)
                    with open(path, 'w') as f:
                        if content:
                            f.write(content)
        except Exception as e:
            self.add_error(f"Failed to create directory structure: {str(e)}")

    def save_file(self, file_content: Union[bytes, Image.Image], 
                 filename: str, directory: Optional[str] = None,
                 category: str = 'misc') -> Optional[str]:
        """Save file with proper organization"""
        try:
            if not directory:
                directory = os.path.join(self.media_dir, category)
            
            os.makedirs(directory, exist_ok=True)
            
            # Generate unique filename
            unique_filename = self.generate_unique_filename(filename, directory)
            file_path = os.path.join(directory, unique_filename)
            
            # Handle different content types
            if isinstance(file_content, Image.Image):
                file_content.save(file_path, quality=95, optimize=True)
            else:
                with open(file_path, 'wb') as f:
                    f.write(file_content)
            
            logger.info(f"File saved successfully: {file_path}")
            return file_path
        except Exception as e:
            self.add_error(f"Failed to save file: {str(e)}")
            return None

    def save_image(self, image: Image.Image, filename: Optional[str] = None) -> Optional[str]:
        """Save image with optimization"""
        try:
            if not filename:
                filename = 'image.jpg'
            return self.save_file(image, filename, 
                                os.path.join(self.media_dir, 'images'))
        except Exception as e:
            self.add_error(f"Failed to save image: {str(e)}")
            return None

    def create_temp_file(self, content: Union[bytes, str], 
                        extension: str = '') -> Optional[str]:
        """Create temporary file with content"""
        try:
            with tempfile.NamedTemporaryFile(
                suffix=extension,
                dir=self.temp_dir,
                delete=False
            ) as temp_file:
                if isinstance(content, str):
                    temp_file.write(content.encode())
                else:
                    temp_file.write(content)
                return temp_file.name
        except Exception as e:
            self.add_error(f"Failed to create temp file: {str(e)}")
            return None

    def clean_temp_files(self, max_age: timedelta = timedelta(hours=24)) -> None:
        """Clean temporary files older than max_age"""
        try:
            current_time = datetime.now()
            for file_name in os.listdir(self.temp_dir):
                file_path = os.path.join(self.temp_dir, file_name)
                if os.path.isfile(file_path):
                    file_time = datetime.fromtimestamp(os.path.getctime(file_path))
                    if current_time - file_time > max_age:
                        os.remove(file_path)
                        logger.info(f"Cleaned temp file: {file_path}")
        except Exception as e:
            self.add_error(f"Failed to clean temp files: {str(e)}")

    def create_zip_archive(self, files: List[str], 
                          output_filename: str,
                          base_dir: Optional[str] = None) -> Optional[str]:
        """Create ZIP archive from files"""
        try:
            if not base_dir:
                base_dir = self.media_dir
                
            archive_path = os.path.join(base_dir, output_filename)
            with zipfile.ZipFile(archive_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for file_path in files:
                    if os.path.exists(file_path):
                        arcname = os.path.basename(file_path)
                        zipf.write(file_path, arcname)
            
            logger.info(f"Created ZIP archive: {archive_path}")
            return archive_path
        except Exception as e:
            self.add_error(f"Failed to create ZIP archive: {str(e)}")
            return None

    def move_file(self, source: str, destination: str, 
                 overwrite: bool = False) -> Optional[str]:
        """Move file with validation"""
        try:
            if not os.path.exists(source):
                raise ValueError(f"Source file does not exist: {source}")
                
            if os.path.exists(destination) and not overwrite:
                raise ValueError(f"Destination file exists: {destination}")
                
            os.makedirs(os.path.dirname(destination), exist_ok=True)
            shutil.move(source, destination)
            
            logger.info(f"Moved file: {source} -> {destination}")
            return destination
        except Exception as e:
            self.add_error(f"Failed to move file: {str(e)}")
            return None

    def copy_file(self, source: str, destination: str, 
                 overwrite: bool = False) -> Optional[str]:
        """Copy file with validation"""
        try:
            if not os.path.exists(source):
                raise ValueError(f"Source file does not exist: {source}")
                
            if os.path.exists(destination) and not overwrite:
                raise ValueError(f"Destination file exists: {destination}")
                
            os.makedirs(os.path.dirname(destination), exist_ok=True)
            shutil.copy2(source, destination)
            
            logger.info(f"Copied file: {source} -> {destination}")
            return destination
        except Exception as e:
            self.add_error(f"Failed to copy file: {str(e)}")
            return None

    def delete_file(self, file_path: str, secure: bool = False) -> bool:
        """Delete file with optional secure deletion"""
        try:
            if not os.path.exists(file_path):
                raise ValueError(f"File does not exist: {file_path}")
                
            if secure:
                # Overwrite file with zeros before deletion
                with open(file_path, 'wb') as f:
                    f.write(b'\0' * os.path.getsize(file_path))
                    
            os.remove(file_path)
            logger.info(f"Deleted file: {file_path}")
            return True
        except Exception as e:
            self.add_error(f"Failed to delete file: {str(e)}")
            return False

    def get_file_info(self, file_path: str) -> Optional[Dict[str, Any]]:
        """Get detailed file information"""
        try:
            if not os.path.exists(file_path):
                raise ValueError(f"File does not exist: {file_path}")
                
            stat = os.stat(file_path)
            return {
                'name': os.path.basename(file_path),
                'path': file_path,
                'size': stat.st_size,
                'created': datetime.fromtimestamp(stat.st_ctime),
                'modified': datetime.fromtimestamp(stat.st_mtime),
                'mime_type': self.get_mime_type(file_path),
                'extension': os.path.splitext(file_path)[1],
            }
        except Exception as e:
            self.add_error(f"Failed to get file info: {str(e)}")
            return None

    def organize_files(self, directory: str, 
                      organize_by: str = 'type') -> Optional[Dict[str, List[str]]]:
        """Organize files by type or date"""
        try:
            organized = {}
            for root, _, files in os.walk(directory):
                for file in files:
                    file_path = os.path.join(root, file)
                    
                    if organize_by == 'type':
                        mime_type = self.get_mime_type(file_path)
                        category = mime_type.split('/')[0] if mime_type else 'other'
                    else:  # organize by date
                        stat = os.stat(file_path)
                        date = datetime.fromtimestamp(stat.st_mtime).strftime('%Y-%m-%d')
                        category = date
                        
                    if category not in organized:
                        organized[category] = []
                    organized[category].append(file_path)
                    
            return organized
        except Exception as e:
            self.add_error(f"Failed to organize files: {str(e)}")
            return None

    def save_document(self, file_content: Union[bytes, BinaryIO], 
                     filename: str, optimize: bool = True) -> Optional[str]:
        """Save document with optimization and proper categorization"""
        try:
            ext = os.path.splitext(filename)[1].lower()
            if ext not in self.DOCUMENT_EXTENSIONS:
                raise ValueError(f"Unsupported document type: {ext}")
            
            # Determine category based on extension
            if ext == '.pdf':
                category = 'documents/pdf'
            elif ext in {'.doc', '.docx'}:
                category = 'documents/word'
            elif ext == '.txt':
                category = 'documents/text'
            else:
                category = 'documents/other'
            
            # Save and optimize if needed
            if optimize:
                if ext == '.pdf':
                    return self._save_optimized_pdf(file_content, filename, category)
                elif ext in {'.doc', '.docx'}:
                    return self._save_optimized_word(file_content, filename, category)
            
            return self.save_file(file_content, filename, category=category)
        except Exception as e:
            self.add_error(f"Failed to save document: {str(e)}")
            return None

    def _save_optimized_pdf(self, file_content: Union[bytes, BinaryIO], 
                          filename: str, category: str) -> Optional[str]:
        """Save PDF with optimization"""
        try:
            # Create temporary file
            temp_path = self.create_temp_file(
                file_content if isinstance(file_content, bytes) else file_content.read(),
                '.pdf'
            )
            
            # Optimize PDF
            pdf_doc = fitz.open(temp_path)
            optimized_path = os.path.join(self.temp_dir, 'optimized.pdf')
            
            # Reduce image quality and compress
            for page in pdf_doc:
                for img in page.get_images():
                    xref = img[0]
                    image = fitz.Pixmap(pdf_doc, xref)
                    if image.n >= 4:  # CMYK
                        image = fitz.Pixmap(fitz.csRGB, image)
                    image.save(os.path.join(self.temp_dir, f"temp_img_{xref}.jpg"), 
                             quality=85)
            
            pdf_doc.save(optimized_path,
                          garbage=4,  # Maximum garbage collection
                          clean=True,  # Clean unused elements
                          deflate=True)  # Compress streams
            pdf_doc.close()
            
            # Save optimized file
            with open(optimized_path, 'rb') as f:
                final_path = self.save_file(f.read(), filename, category=category)
            
            # Cleanup
            os.remove(temp_path)
            os.remove(optimized_path)
            
            return final_path
        except Exception as e:
            self.add_error(f"Failed to optimize PDF: {str(e)}")
            return None

    def _save_optimized_word(self, file_content: Union[bytes, BinaryIO], 
                           filename: str, category: str) -> Optional[str]:
        """Save Word document with optimization"""
        try:
            # Create temporary file
            temp_path = self.create_temp_file(
                file_content if isinstance(file_content, bytes) else file_content.read(),
                '.docx'
            )
            
            # Open and optimize document
            doc = Document(temp_path)
            
            # Optimize images in document
            for shape in doc.inline_shapes:
                if shape.width > Inches(6):  # Resize large images
                    shape.width = Inches(6)
                    shape.height = shape.width * shape.aspect_ratio
            
            # Save optimized document
            optimized_path = os.path.join(self.temp_dir, 'optimized.docx')
            doc.save(optimized_path)
            
            # Save final file
            with open(optimized_path, 'rb') as f:
                final_path = self.save_file(f.read(), filename, category=category)
            
            # Cleanup
            os.remove(temp_path)
            os.remove(optimized_path)
            
            return final_path
        except Exception as e:
            self.add_error(f"Failed to optimize Word document: {str(e)}")
            return None

    def extract_document_text(self, file_path: str) -> Optional[str]:
        """Extract text content from document"""
        try:
            ext = os.path.splitext(file_path)[1].lower()
            
            if ext == '.pdf':
                with fitz.open(file_path) as doc:
                    text = ""
                    for page in doc:
                        text += page.get_text()
                    return text
            elif ext in {'.doc', '.docx'}:
                doc = Document(file_path)
                return '\n'.join(paragraph.text for paragraph in doc.paragraphs)
            elif ext == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            else:
                raise ValueError(f"Unsupported document type for text extraction: {ext}")
        except Exception as e:
            self.add_error(f"Failed to extract document text: {str(e)}")
            return None

    def get_document_metadata(self, file_path: str) -> Optional[Dict[str, Any]]:
        """Get detailed document metadata"""
        try:
            ext = os.path.splitext(file_path)[1].lower()
            base_info = self.get_file_info(file_path)
            
            if not base_info:
                return None
            
            metadata = {**base_info}
            
            if ext == '.pdf':
                with fitz.open(file_path) as doc:
                    metadata.update({
                        'page_count': len(doc),
                        'title': doc.metadata.get('title', ''),
                        'author': doc.metadata.get('author', ''),
                        'subject': doc.metadata.get('subject', ''),
                        'keywords': doc.metadata.get('keywords', ''),
                        'page_size': doc[0].rect.size,
                    })
            elif ext in {'.doc', '.docx'}:
                doc = Document(file_path)
                core_props = doc.core_properties
                metadata.update({
                    'title': core_props.title or '',
                    'author': core_props.author or '',
                    'created': core_props.created,
                    'modified': core_props.modified,
                    'last_modified_by': core_props.last_modified_by or '',
                    'revision': core_props.revision,
                    'paragraph_count': len(doc.paragraphs),
                })
            
            return metadata
        except Exception as e:
            self.add_error(f"Failed to get document metadata: {str(e)}")
            return None

    def convert_document(self, source_path: str, target_format: str) -> Optional[str]:
        """Convert document to different format"""
        try:
            source_ext = os.path.splitext(source_path)[1].lower()
            if source_ext not in self.DOCUMENT_EXTENSIONS:
                raise ValueError(f"Unsupported source format: {source_ext}")
            
            if target_format not in {'.pdf', '.docx', '.txt'}:
                raise ValueError(f"Unsupported target format: {target_format}")
            
            # Generate output filename
            output_name = os.path.splitext(os.path.basename(source_path))[0]
            output_path = os.path.join(
                self.media_dir, 
                'documents/converted', 
                f"{output_name}{target_format}"
            )
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            if source_ext == '.pdf' and target_format == '.txt':
                # PDF to Text
                text = self.extract_document_text(source_path)
                if text:
                    with open(output_path, 'w', encoding='utf-8') as f:
                        f.write(text)
            elif source_ext == '.pdf' and target_format == '.docx':
                # PDF to Word
                doc = Document()
                text = self.extract_document_text(source_path)
                if text:
                    doc.add_paragraph(text)
                    doc.save(output_path)
            elif source_ext in {'.doc', '.docx'} and target_format == '.pdf':
                # Word to PDF
                doc = fitz.open()
                text = self.extract_document_text(source_path)
                if text:
                    page = doc.new_page()
                    page.insert_text((72, 72), text)
                    doc.save(output_path)
                doc.close()
            
            return output_path if os.path.exists(output_path) else None
        except Exception as e:
            self.add_error(f"Failed to convert document: {str(e)}")
            return None

    def merge_pdfs(self, pdf_files: List[str], output_filename: str) -> Optional[str]:
        """Merge multiple PDF files into one"""
        try:
            output_path = os.path.join(self.media_dir, 'documents/merged', output_filename)
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            merged_pdf = fitz.open()
            
            for pdf_file in pdf_files:
                if not self.validate_file_type(pdf_file, {'.pdf'}):
                    raise ValueError(f"Invalid PDF file: {pdf_file}")
                    
                pdf_doc = fitz.open(pdf_file)
                merged_pdf.insert_pdf(pdf_doc)
                pdf_doc.close()
            
            merged_pdf.save(output_path,
                          garbage=4,  # Maximum garbage collection
                          clean=True,  # Clean unused elements
                          deflate=True)  # Compress streams
            merged_pdf.close()
            
            logger.info(f"Successfully merged PDFs into: {output_path}")
            return output_path
        except Exception as e:
            self.add_error(f"Failed to merge PDFs: {str(e)}")
            return None

    def split_pdf(self, pdf_file: str, output_dir: Optional[str] = None) -> Optional[List[str]]:
        """Split PDF into individual pages"""
        try:
            if not output_dir:
                output_dir = os.path.join(self.media_dir, 'documents/split')
            os.makedirs(output_dir, exist_ok=True)
            
            if not self.validate_file_type(pdf_file, {'.pdf'}):
                raise ValueError(f"Invalid PDF file: {pdf_file}")
            
            pdf_doc = fitz.open(pdf_file)
            output_files = []
            
            for page_num in range(len(pdf_doc)):
                # Create new PDF with single page
                output_pdf = fitz.open()
                output_pdf.insert_pdf(pdf_doc, from_page=page_num, to_page=page_num)
                
                # Save page
                output_path = os.path.join(output_dir, f"page_{page_num + 1}.pdf")
                output_pdf.save(output_path,
                              garbage=4,
                              clean=True,
                              deflate=True)
                output_pdf.close()
                output_files.append(output_path)
            
            pdf_doc.close()
            logger.info(f"Successfully split PDF into {len(output_files)} pages")
            return output_files
        except Exception as e:
            self.add_error(f"Failed to split PDF: {str(e)}")
            return None

    def add_pdf_watermark(self, pdf_file: str, watermark_text: str, 
                         output_filename: Optional[str] = None) -> Optional[str]:
        """Add watermark to PDF"""
        try:
            if not output_filename:
                output_filename = os.path.basename(pdf_file).replace('.pdf', '_watermarked.pdf')
            output_path = os.path.join(self.media_dir, 'documents/watermarked', output_filename)
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            pdf_doc = fitz.open(pdf_file)
            
            for page in pdf_doc:
                # Calculate center position
                rect = page.rect
                text_width = len(watermark_text) * 12  # Approximate width
                x = (rect.width - text_width) / 2
                y = rect.height / 2
                
                # Add watermark
                page.insert_text((x, y),
                               watermark_text,
                               fontsize=50,
                               rotate=45,
                               color=(0.7, 0.7, 0.7),  # Light gray
                               opacity=0.3)
            
            pdf_doc.save(output_path,
                        garbage=4,
                        clean=True,
                        deflate=True)
            pdf_doc.close()
            
            logger.info(f"Successfully added watermark to PDF: {output_path}")
            return output_path
        except Exception as e:
            self.add_error(f"Failed to add PDF watermark: {str(e)}")
            return None

    def encrypt_pdf(self, pdf_file: str, user_password: str, owner_password: Optional[str] = None,
                   output_filename: Optional[str] = None) -> Optional[str]:
        """Encrypt PDF with password protection"""
        try:
            if not output_filename:
                output_filename = os.path.basename(pdf_file).replace('.pdf', '_encrypted.pdf')
            output_path = os.path.join(self.media_dir, 'documents/encrypted', output_filename)
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            pdf_doc = fitz.open(pdf_file)
            
            # Set passwords and permissions
            pdf_doc.save(output_path,
                        garbage=4,
                        clean=True,
                        deflate=True,
                        encryption=fitz.PDF_ENCRYPT_AES_256,  # Strong encryption
                        user_pw=user_password,
                        owner_pw=owner_password or user_password,
                        permissions=int(
                            fitz.PDF_PERM_PRINT |  # Allow printing
                            fitz.PDF_PERM_COPY |   # Allow copying
                            fitz.PDF_PERM_READ     # Allow reading
                        ))
            pdf_doc.close()
            
            logger.info(f"Successfully encrypted PDF: {output_path}")
            return output_path
        except Exception as e:
            self.add_error(f"Failed to encrypt PDF: {str(e)}")
            return None

    def compress_pdf(self, pdf_file: str, quality: str = 'medium',
                    output_filename: Optional[str] = None) -> Optional[str]:
        """Compress PDF with different quality settings"""
        try:
            if not output_filename:
                output_filename = os.path.basename(pdf_file).replace('.pdf', '_compressed.pdf')
            output_path = os.path.join(self.media_dir, 'documents/compressed', output_filename)
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            # Quality settings
            quality_settings = {
                'low': {'image_quality': 50, 'dpi': 72},
                'medium': {'image_quality': 75, 'dpi': 150},
                'high': {'image_quality': 90, 'dpi': 300}
            }
            settings = quality_settings.get(quality, quality_settings['medium'])
            
            pdf_doc = fitz.open(pdf_file)
            
            # Process each page
            for page in pdf_doc:
                # Compress images on the page
                for img in page.get_images():
                    xref = img[0]
                    image = fitz.Pixmap(pdf_doc, xref)
                    
                    if image.n >= 4:  # CMYK
                        image = fitz.Pixmap(fitz.csRGB, image)
                    
                    # Reduce image quality
                    image.set_dpi(settings['dpi'], settings['dpi'])
                    image.save(
                        os.path.join(self.temp_dir, f"temp_img_{xref}.jpg"),
                        quality=settings['image_quality']
                    )
            
            # Save compressed PDF
            pdf_doc.save(output_path,
                        garbage=4,
                        clean=True,
                        deflate=True)
            pdf_doc.close()
            
            logger.info(f"Successfully compressed PDF: {output_path}")
            return output_path
        except Exception as e:
            self.add_error(f"Failed to compress PDF: {str(e)}")
            return None

    def extract_zip_archive(self, archive_path: str, 
                          extract_dir: Optional[str] = None) -> Optional[Dict]:
        """Extract ZIP archive with metadata handling"""
        try:
            if not extract_dir:
                extract_dir = os.path.join(
                    self.temp_dir,
                    f"extract_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
                )
            
            os.makedirs(extract_dir, exist_ok=True)
            
            extracted_files = []
            metadata = {}
            
            with zipfile.ZipFile(archive_path, 'r') as zipf:
                # Extract metadata if present
                if 'metadata.json' in zipf.namelist():
                    with zipf.open('metadata.json') as meta_file:
                        metadata = json.load(meta_file)
                
                # Extract all files
                for file_info in zipf.filelist:
                    if file_info.filename != 'metadata.json':
                        extracted_path = zipf.extract(file_info, extract_dir)
                        extracted_files.append(extracted_path)
            
            return {
                'files': extracted_files,
                'metadata': metadata,
                'extract_dir': extract_dir
            }
        except Exception as e:
            self.add_error(f"Failed to extract ZIP archive: {str(e)}")
            return None

    def save_optimized_image(self, image_data: Union[bytes, BinaryIO, Image.Image],
                           max_size: int = 1024) -> Optional[str]:
        """Save image with optimization"""
        try:
            if isinstance(image_data, (bytes, BinaryIO)):
                image = Image.open(image_data)
            elif isinstance(image_data, Image.Image):
                image = image_data
            else:
                raise ValueError("Invalid image data type")
            
            # Resize if needed
            if max(image.size) > max_size:
                ratio = max_size / max(image.size)
                new_size = tuple(int(dim * ratio) for dim in image.size)
                image = image.resize(new_size, Image.Resampling.LANCZOS)
            
            # Convert to RGB if needed
            if image.mode in ('RGBA', 'P'):
                image = image.convert('RGB')
            
            # Generate filename
            filename = self.generate_unique_filename('image.jpg')
            file_path = os.path.join(self.media_dir, 'images', filename)
            
            # Save with optimization
            image.save(file_path, 'JPEG', quality=85, optimize=True)
            
            return file_path
        except Exception as e:
            self.add_error(f"Failed to save optimized image: {str(e)}")
            return None

    def save_optimized_video(self, video_data: BinaryIO,
                           max_size: int = 1920) -> Optional[str]:
        """Save video with optimization"""
        try:
            # Save original video first
            temp_path = self.create_temp_file(video_data.read(), '.mp4')
            if not temp_path:
                raise ValueError("Failed to save temporary video")
            
            # Generate output filename
            filename = self.generate_unique_filename('video.mp4')
            output_path = os.path.join(self.media_dir, 'videos', filename)
            
            # Use ffmpeg for optimization
            os.system(f'ffmpeg -i {temp_path} -vf "scale=min({max_size}\\:iw):-2" '
                     f'-c:v libx264 -crf 23 -preset medium -c:a aac -b:a 128k '
                     f'{output_path}')
            
            # Clean up temp file
            os.remove(temp_path)
            
            return output_path
        except Exception as e:
            self.add_error(f"Failed to save optimized video: {str(e)}")
            return None

    def save_optimized_audio(self, audio_data: BinaryIO,
                           format: str = 'mp3',
                           bitrate: str = '192k') -> Optional[str]:
        """Save audio with optimization"""
        try:
            # Save original audio first
            temp_path = self.create_temp_file(audio_data.read(), f'.{format}')
            if not temp_path:
                raise ValueError("Failed to save temporary audio")
            
            # Generate output filename
            filename = self.generate_unique_filename(f'audio.{format}')
            output_path = os.path.join(self.media_dir, 'audio', filename)
            
            # Use ffmpeg for optimization
            os.system(f'ffmpeg -i {temp_path} -c:a libmp3lame -b:a {bitrate} '
                     f'{output_path}')
            
            # Clean up temp file
            os.remove(temp_path)
            
            return output_path
        except Exception as e:
            self.add_error(f"Failed to save optimized audio: {str(e)}")
            return None

    def monitor_file_changes(self, directory: str,
                           callback: Optional[Callable] = None) -> None:
        """Monitor directory for file changes"""
        try:
            import watchdog.observers
            import watchdog.events
            
            class FileChangeHandler(watchdog.events.FileSystemEventHandler):
                def __init__(self, callback):
                    self.callback = callback
                
                def on_any_event(self, event):
                    if self.callback:
                        self.callback(event)
            
            observer = watchdog.observers.Observer()
            observer.schedule(
                FileChangeHandler(callback),
                directory,
                recursive=True
            )
            observer.start()
            
            logger.info(f"Started monitoring directory: {directory}")
            
        except Exception as e:
            self.add_error(f"Failed to start file monitoring: {str(e)}")

    def get_file_version(self, file_path: str) -> Optional[str]:
        """Get file version information"""
        try:
            stat = os.stat(file_path)
            return {
                'path': file_path,
                'size': stat.st_size,
                'created': datetime.fromtimestamp(stat.st_ctime),
                'modified': datetime.fromtimestamp(stat.st_mtime),
                'accessed': datetime.fromtimestamp(stat.st_atime),
                'checksum': self._calculate_checksum(file_path)
            }
        except Exception as e:
            self.add_error(f"Failed to get file version: {str(e)}")
            return None

    def _calculate_checksum(self, file_path: str) -> str:
        """Calculate file checksum"""
        try:
            hasher = hashlib.md5()
            with open(file_path, 'rb') as f:
                for chunk in iter(lambda: f.read(4096), b''):
                    hasher.update(chunk)
            return hasher.hexdigest()
        except Exception as e:
            self.add_error(f"Failed to calculate checksum: {str(e)}")
            return ''